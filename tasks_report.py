import comtypes.client
from docx import Document
import os
import mysql.connector
from datetime import datetime

def get_versioned_filename(base_name, extension, directory):
    """Generate a filename with an incremented version number in the specified directory."""
    counter = 1
    while True:
        filename = f"{base_name}_v{counter}.{extension}"
        full_path = os.path.join(directory, filename)
        if not os.path.exists(full_path):
            return full_path
        counter += 1

def replace_placeholders(doc_path, output_path, data):
    """Replace placeholders in the Word document with dynamic data."""
    if not os.path.exists(doc_path):
        print(f"Template file not found at: {doc_path}")
        exit(1)

    doc = Document(doc_path)

    # Debug: Print all paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: {paragraph.text}")

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                print(f"Replacing {key} with {value}")
                paragraph.text = paragraph.text.replace(key, value)

    # Save the updated Word document
    doc.save(output_path)

def convert_to_pdf(word_path, pdf_path):
    """Convert the Word document to a PDF."""
    word = comtypes.client.CreateObject("Word.Application")
    word.Visible = False  # Hide Word application
    try:
        doc = word.Documents.Open(word_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 represents wdFormatPDF
        doc.Close()
    finally:
        word.Quit()

def fetch_data_from_database(user_id):
    """Fetch data from the database for a specific user and prepare it for the template."""
    try:
        # Database connection
        db = mysql.connector.connect(
            host="localhost",
            user="root",
            password="",
            database="login_system"
        )
        cursor = db.cursor()

        # Fetch total users
        cursor.execute("SELECT COUNT(*) FROM users")
        total_users = cursor.fetchone()[0]

        # Fetch total tasks
        cursor.execute("SELECT COUNT(*) FROM tasks")
        total_tasks = cursor.fetchone()[0]

        # Fetch user-wise statistics for the specified user
        cursor.execute("""
            SELECT
                u.id AS user_id,
                u.username,
                u.department,
                u.role,
                u.email,
                COUNT(t.task_id) AS total_tasks_assigned,
                SUM(CASE WHEN t.status = 'Pending' THEN 1 ELSE 0 END) AS tasks_pending,
                SUM(CASE WHEN t.status = 'Started' THEN 1 ELSE 0 END) AS tasks_started,
                SUM(CASE WHEN t.status = 'Completed on Time' THEN 1 ELSE 0 END) AS tasks_completed_on_time,
                SUM(CASE WHEN t.status = 'Delayed Completion' THEN 1 ELSE 0 END) AS tasks_delayed,
                SUM(CASE WHEN t.status IN ('Completed on Time', 'Delayed Completion') THEN 1 ELSE 0 END) AS tasks_completed_total,
                SUM(CASE WHEN t.status IN ('Pending', 'Started') THEN 1 ELSE 0 END) AS tasks_not_completed,
                ROUND(100.0 * SUM(CASE WHEN t.status = 'Completed on Time' THEN 1 ELSE 0 END) / COUNT(t.task_id), 2) AS on_time_completion_rate,
                ROUND(100.0 * SUM(CASE WHEN t.status IN ('Completed on Time', 'Delayed Completion') THEN 1 ELSE 0 END) / COUNT(t.task_id), 2) AS overall_completion_rate,
                AVG(DATEDIFF(t.actual_completion_date, t.expected_finish_date)) AS average_delay_duration
            FROM users u
            LEFT JOIN tasks t ON u.id = t.user_id
            WHERE u.id = %s
            GROUP BY u.id
        """, (user_id,))
        user = cursor.fetchone()

        # Close the database connection
        cursor.close()
        db.close()

        if user:
            user_data = {
                "username": user[1],
                "id": user[0],
                "department": user[2],
                "role": user[3],
                "email": user[4],
                "total_tasks_assigned": user[5],
                "tasks_pending": user[6],
                "tasks_started": user[7],
                "tasks_completed_on_time": user[8],
                "tasks_delayed": user[9],
                "tasks_completed_total": user[10],
                "tasks_not_completed": user[11],
                "on_time_completion_rate": user[12],
                "overall_completion_rate": user[13],
                "average_delay_duration": user[14]
            }
        else:
            user_data = None

        return {
            "total_users": total_users,
            "total_tasks": total_tasks,
            "user": user_data
        }

    except mysql.connector.Error as err:
        print(f"Database error: {err}")
        return None

if __name__ == '__main__':
    # Prompt the user to enter the user_id
    user_id = input("Enter the user ID: ")

    # Fetch data from the database for the specified user
    data = fetch_data_from_database(user_id)

    if data is None:
        print("Failed to fetch data from the database.")
    elif data["user"] is None:
        print(f"No user found with ID: {user_id}")
    else:
        # Prepare the template data
        template_data = {
            "{{ total_users }}": str(data["total_users"]),
            "{{ total_tasks }}": str(data["total_tasks"]),
            "{{ user.username }}": data["user"]["username"],
            "{{ user.id }}": str(data["user"]["id"]),
            "{{ user.department }}": data["user"]["department"],
            "{{ user.role }}": data["user"]["role"],
            "{{ user.email }}": data["user"]["email"],
            "{{ user.total_tasks_assigned }}": str(data["user"]["total_tasks_assigned"]),
            "{{ user.tasks_pending }}": str(data["user"]["tasks_pending"]),
            "{{ user.tasks_started }}": str(data["user"]["tasks_started"]),
            "{{ user.tasks_completed_on_time }}": str(data["user"]["tasks_completed_on_time"]),
            "{{ user.tasks_delayed }}": str(data["user"]["tasks_delayed"]),
            "{{ user.tasks_completed_total }}": str(data["user"]["tasks_completed_total"]),
            "{{ user.tasks_not_completed }}": str(data["user"]["tasks_not_completed"]),
            "{{ user.on_time_completion_rate }}": str(data["user"]["on_time_completion_rate"]),
            "{{ user.overall_completion_rate }}": str(data["user"]["overall_completion_rate"]),
            "{{ user.average_delay_duration }}": str(data["user"]["average_delay_duration"])
        }

        # Get absolute paths
        base_dir = os.path.dirname(os.path.abspath(__file__))
        reports_dir = os.path.join(base_dir, 'reports')
        template_path = os.path.join(base_dir, 'tasks_report.docx')

        # Ensure the Reports directory exists
        os.makedirs(reports_dir, exist_ok=True)

        # Generate versioned filenames in the Reports directory
        word_output_path = get_versioned_filename(f"user_{user_id}_task_report", "docx", reports_dir)
        pdf_output_path = get_versioned_filename(f"user_{user_id}_task_report", "pdf", reports_dir)

        # Process and generate the outputs
        replace_placeholders(template_path, word_output_path, template_data)
        convert_to_pdf(word_output_path, pdf_output_path)

        print(f'Report generated:\nWord file: {word_output_path}\nPDF file: {pdf_output_path}')